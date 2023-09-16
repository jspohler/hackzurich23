import sqlite3
import os
import glob
from docx import Document
import extract_msg
import fitz 
import olefile
import markdown
import re
import ast
from bs4 import BeautifulSoup


def extract_data_from_db(db_file, output_folder):
    # Extract data from a single SQLite database file

    conn = sqlite3.connect(db_file)
    cursor = conn.cursor()
    
    # Get the list of table names from the database schema
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
    table_names = [table[0] for table in cursor.fetchall()]

    # Create a data file for the database and its tables
    db_file_name = os.path.splitext(os.path.basename(db_file))[0]
    txt_file_path = os.path.join(output_folder, f"{db_file_name}_db.txt")

    with open(txt_file_path, 'w') as txt_file:
        txt_file.write(f"{db_file_name}\n")

        # Loop through each table and fetch text
        for table_name in table_names:
            cursor.execute(f"SELECT * FROM {table_name};")
            rows = cursor.fetchall()

            txt_file.write(f"{table_name}\n")
            for row in rows:
                txt_file.write('\t'.join(map(str, row[1:])) + '\n')
    # Close the database connection
    conn.close()

def extract_db_data(folder_path, output_folder):
    # Extract data from all SQLite database files in the specified folder
    os.makedirs(output_folder, exist_ok=True)

    db_files = glob.glob(os.path.join(folder_path, '*.db'))

    for db_file in db_files:
        extract_data_from_db(db_file, output_folder)
    
def extract_data_from_docx(docx_file, output_folder):
    # Extract data from a single docx file
    doc = Document(docx_file)
    
    # Create a text file for the .docx file
    txt_file_path = os.path.join(output_folder, f"{os.path.splitext(os.path.basename(docx_file))[0]}_docx.txt")
    with open(txt_file_path, 'w', encoding='utf-8') as txt_file:
        for paragraph in doc.paragraphs:
            txt_file.write(paragraph.text + '\n')

def extract_str_from_docx(docx_file):
    doc = Document(docx_file)
    full_text = []
    
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    
    return '\n'.join(full_text)

def extract_docx_data(folder_path, output_folder):
    # Extract data from all docx files in the specified folder
    os.makedirs(output_folder, exist_ok=True)
    docx_files = glob.glob(os.path.join(folder_path, '*.docx'))
    
    for docx_file in docx_files:
        extract_data_from_docx(docx_file, output_folder)

def copy_data_from_log(log_file, output_folder):
    #TODO when doing tokenizing it may require changes
    # Extract data from a single log file
    base_name = os.path.splitext(os.path.basename(log_file))[0]

    # Rename the .log file to .txt
    txt_file_path = os.path.join(output_folder, f"{base_name}_log.txt")
    with open(log_file, 'r', encoding='ISO-8859-1') as log_file_content:
        with open(txt_file_path, 'w', encoding='utf-8') as txt_file:
            txt_file.write(log_file_content.read())

def extract_log_data(folder_path, output_folder):
    # Extract data from all log files in the specified folder
    os.makedirs(output_folder, exist_ok=True)
    log_files = glob.glob(os.path.join(folder_path, '*.log'))
    
    for log_file in log_files:
        copy_data_from_log(log_file, output_folder)

def copy_data_from_md(md_file, output_folder):
    
    with open(md_file, 'r', encoding='ISO-8859-1') as f:
        text = f.read()
        text = text.replace("```", "")
        html = markdown.markdown(text)
        # extract text
        soup = BeautifulSoup(html, "lxml")
        text = ''.join(soup.findAll(text=True))

        
        file_name = os.path.splitext(os.path.basename(md_file))[0]
        txt_file_path = os.path.join(output_folder, f"{file_name}_md.txt")

        with open(txt_file_path, 'w', encoding='utf-8', errors='ignore') as txt_file:
            txt_file.write(text)

def extract_md_data(folder_path, output_folder):
    # Extract data from all md files in the specified folder
    os.makedirs(output_folder, exist_ok=True)
    md_files = glob.glob(os.path.join(folder_path, '*.md'))
    
    for md_file in md_files:
        copy_data_from_md(md_file, output_folder)

def extract_data_from_msg(msg_file, output_folder):
    base_name = os.path.splitext(os.path.basename(msg_file))[0]
    txt_file_path = os.path.join(output_folder, f"{base_name}_msg.txt")
    
    try:
        msg = extract_msg.Message(msg_file)
        sender = msg.sender if hasattr(msg, 'sender') else "N/A"
        recipients = msg.to if hasattr(msg, 'to') else "N/A"
        subject = msg.subject if hasattr(msg, 'subject') else "N/A"
        email_text = msg.body if hasattr(msg, 'body') else "N/A"
    except (UnicodeEncodeError, AttributeError, TypeError):
        pass
    
    with open(txt_file_path, 'w', encoding='utf-8') as txt_file:
        txt_file.write(f"Sender: {sender}\n")
        txt_file.write(f"Recipients: {recipients}\n")
        txt_file.write(f"Subject: {subject}\n")
        txt_file.write("Email Text:\n")
        txt_file.write(email_text)

def extract_msg_data(folder_path, output_folder):
    # Extract data from all msg files in the specified folder
    os.makedirs(output_folder, exist_ok=True)
    msg_files = glob.glob(os.path.join(folder_path, '*.msg'))
    
    for msg_file in msg_files:
        extract_data_from_msg(msg_file, output_folder)

def copy_data_from_ps1(ps1_file, output_folder):
    # Get the base name of the file (without extension)
    base_name = os.path.splitext(os.path.basename(ps1_file))[0]

    # Create a new .txt file in the output folder and copy the content
    txt_file_path = os.path.join(output_folder, f"{base_name}_ps1.txt")
    with open(ps1_file, 'r', encoding='ISO-8859-1') as ps1_file_content:
        with open(txt_file_path, 'w', encoding='utf-8') as txt_file:
            txt_file.write(ps1_file_content.read())

def extract_ps1_data(folder_path, output_folder):
    os.makedirs(output_folder, exist_ok=True)

    ps1_files = glob.glob(os.path.join(folder_path, '*.ps1'))

    for ps1_file in ps1_files:
        copy_data_from_ps1(ps1_file, output_folder)
    
#This code will attempt to extract the text from the files but may not preserve the layout perfectly.
def extract_data_from_pub(pub_file, output_folder):
    base_name = os.path.splitext(os.path.basename(pub_file))[0]

    txt_file_path = os.path.join(output_folder, f"{base_name}_pub.txt")
    try:
        with open(pub_file, 'r', encoding='utf-8', errors='ignore') as pub_file_content:
            rsa_key_text = pub_file_content.read()
        
        # Save the extracted RSA key text to a .txt file
        with open(txt_file_path, 'w', encoding='utf-8') as txt_file:
            txt_file.write(rsa_key_text)
        
    except Exception as e:
        print(f"Error extracting RSA key from {pub_file}: {e}")

def extract_pub_data(folder_path, output_folder):
    os.makedirs(output_folder, exist_ok=True)
    
    pub_files = glob.glob(os.path.join(folder_path, '*.pub'))
    
    for pub_file in pub_files:
        if pub_file.endswith('.pub'):
            extract_data_from_pub(pub_file, output_folder)

def extract_data_from_pem(pem_file, output_folder):
    base_name = os.path.splitext(os.path.basename(pem_file))[0]

    txt_file_path = os.path.join(output_folder, f"{base_name}_pem.txt")
    try:
        with open(pem_file, 'r', encoding='utf-8', errors='ignore') as pem_file_content:
            pem_data = pem_file_content.read()
        # Save the extracted PEM data to a .txt file
        with open(txt_file_path, 'w', encoding='utf-8') as txt_file:
            txt_file.write(pem_data)
        
    except Exception as e:
        print(f"Error extracting data from {pem_file}: {e}")

def extract_pem_data(folder_path, output_folder):
    os.makedirs(output_folder, exist_ok=True)
    pem_files = glob.glob(os.path.join(folder_path, '*.pem'))
    
    for pem_file in pem_files:
        if pem_file.endswith('.pem'):
            extract_data_from_pem(pem_file, output_folder)

def extract_data_from_py_old(py_file, output_folder):
    base_name = os.path.splitext(os.path.basename(py_file))[0]

    txt_file_path = os.path.join(output_folder, f"{base_name}_py.txt")
    try:
        with open(py_file, 'r', encoding='utf-8') as py_file_content:
            python_code = py_file_content.read()
        with open(txt_file_path, 'w', encoding='utf-8') as txt_file:
            txt_file.write(python_code)
    except Exception as e:
        print(f"Error transforming Python code from {py_file}: {e}")
                    
def extract_py_data(folder_path, output_folder):
    os.makedirs(output_folder, exist_ok=True)
    
    py_files = glob.glob(os.path.join(folder_path, '*.py'))
    
    for py_file in py_files:
        if py_file.endswith('.py'):
            extract_data_from_pub(py_file, output_folder)

def filter_code(code):
    filtered_parts = []

    # Collecting comments and inline comments
    for line in code.split("\n"):
        if "#" in line:
            comment = line.split("#", 1)[1].strip()
            filtered_parts.append(comment)
            line = line.split("#")[0]  # remove the comment part from line

    # Parsing the code into an AST
    try:
        tree = ast.parse(code)
    except SyntaxError:
        print("Couldn't parse the entire source code, proceeding with what's parsable.")
        return filtered_parts

    for node in ast.walk(tree):
        if isinstance(node, ast.Str):  # For string literals
            filtered_parts.append(node.s)
        elif isinstance(node, ast.Num):  # For numbers
            filtered_parts.append(node.n)
        elif isinstance(node, ast.Dict):  # For dictionary literals
            try:
                filtered_parts.append({k.value: v.value for k, v in zip(node.keys, node.values)})
            except:
                print("Dict")
        elif isinstance(node, ast.List):  # For list literals
            try:
                filtered_parts.append([elt.value for elt in node.elts])
            except:
                print("List")
        elif isinstance(node, ast.Tuple):  # For tuple literals
            try:
                filtered_parts.append(tuple(elt.value for elt in node.elts))
            except:
                print("Tuple")
        # Add more conditions here if you need to keep more types of values

    return filtered_parts

def extract_data_from_py(py_file, output_folder):
    
    with open(py_file, "r") as f:
        source_code = f.read()

    # Run the filter
    filtered_parts = filter_code(source_code)
    filtered_text = ' '.join(map(str, filtered_parts))


    os.makedirs(output_folder, exist_ok=True)

    file_name = os.path.splitext(os.path.basename(py_file))[0]
    txt_file_path = os.path.join(output_folder, f"{file_name}_py.txt")

    with open(txt_file_path, 'w', encoding='utf-8', errors='ignore') as txt_file:
        txt_file.write(str(filtered_text))

folder_path = 'files'

#py
if False:
    output_folder_py = 'texted_from_files/py/'
    extract_py_data(folder_path, output_folder_py)

#pem
if False:
    output_folder_pem = 'texted_from_files/pem/'
    extract_pub_data(folder_path, output_folder_pem)
    
#pub
if False:
    output_folder_pub = 'texted_from_files/pub/'
    extract_pub_data(folder_path, output_folder_pub)
        
#ps1
if False:
    output_folder_ps1 = 'texted_from_files/ps1/'
    extract_ps1_data(folder_path, output_folder_ps1)
    
#msg
if False:
    output_folder_msg = 'texted_from_files/msg/'
    extract_msg_data(folder_path, output_folder_msg)
    
#md
if False:
    output_folder_md = 'texted_from_files/md/'
    extract_md_data(folder_path, output_folder_md)
    
#log
if False:
    output_folder_log = 'texted_from_files/log/'
    extract_log_data(folder_path, output_folder_log)

#docx
if False:
    output_folder_docx = 'texted_from_files/docx/'
    extract_docx_data(folder_path, output_folder_docx)
    
#db
if False:
    output_folder_db = 'texted_from_files/db/'
    extract_db_data(folder_path, output_folder_db)

